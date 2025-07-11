import logging
from docx import Document as DocxDocument
from docx.shared import Inches
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Class to handle Word document processing and text extraction"""
    
    def __init__(self):
        self.page_break_indicators = [
            '📖 Page',  # Based on the uploaded document format
            '\f',  # Form feed character
            '\x0c'  # Another form feed representation
        ]
    
    def extract_text_from_docx(self, file_path):
        """
        Extract text content from a Word document page by page
        
        Args:
            file_path (str): Path to the .docx file
            
        Returns:
            list: List of tuples (page_number, content)
        """
        try:
            doc = DocxDocument(file_path)
            pages = []
            current_page = 1
            current_content = []
            
            logger.debug(f"Processing document with {len(doc.paragraphs)} paragraphs")
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Check if this paragraph indicates a new page
                if self._is_page_break(text):
                    # Save current page content if it exists
                    if current_content:
                        page_content = '\n'.join(current_content).strip()
                        if page_content:
                            pages.append((current_page, page_content))
                        current_content = []
                    
                    # Extract page number from the page indicator
                    page_num = self._extract_page_number(text)
                    if page_num:
                        current_page = page_num
                    else:
                        current_page += 1
                    
                    continue
                
                # Add content to current page
                if text:
                    current_content.append(text)
            
            # Add the last page if it has content
            if current_content:
                page_content = '\n'.join(current_content).strip()
                if page_content:
                    pages.append((current_page, page_content))
            
            # If no page breaks were found, treat the entire document as one page
            if not pages:
                all_content = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        all_content.append(text)
                
                if all_content:
                    pages.append((1, '\n'.join(all_content)))
            
            logger.info(f"Extracted {len(pages)} pages from document")
            
            # Process tables if any
            self._extract_tables(doc, pages)
            
            return pages
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _is_page_break(self, text):
        """Check if the text indicates a page break"""
        for indicator in self.page_break_indicators:
            if indicator in text:
                return True
        return False
    
    def _extract_page_number(self, text):
        """Extract page number from page break text"""
        # Look for patterns like "📖 Page 1", "Page 2", etc.
        match = re.search(r'Page\s+(\d+)', text, re.IGNORECASE)
        if match:
            return int(match.group(1))
        return None
    
    def _extract_tables(self, doc, pages):
        """Extract table content and append to relevant pages"""
        try:
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(' | '.join(row_content))
                
                if table_content:
                    table_text = '\n'.join(table_content)
                    # For now, append table content to the last page
                    # In a more sophisticated implementation, we'd track table positions
                    if pages:
                        last_page_num, last_content = pages[-1]
                        pages[-1] = (last_page_num, last_content + '\n\nTable Content:\n' + table_text)
                        
        except Exception as e:
            logger.warning(f"Error extracting tables: {str(e)}")
    
    def count_words(self, text):
        """Count words in the given text"""
        words = re.findall(r'\b\w+\b', text)
        return len(words)
    
    def extract_document_metadata(self, file_path):
        """Extract basic metadata from the document"""
        try:
            doc = DocxDocument(file_path)
            
            # Try to extract title from the first few paragraphs
            title = "Untitled Document"
            subject = "General"
            
            for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
                text = paragraph.text.strip()
                if text and len(text) > 10:
                    # Look for chapter or lesson indicators
                    if any(keyword in text.upper() for keyword in ['CHAPTER', 'LESSON', 'UNIT']):
                        title = text
                        break
                    elif not title or title == "Untitled Document":
                        title = text[:100] + "..." if len(text) > 100 else text
            
            # Try to determine subject from content
            content_text = ' '.join([p.text for p in doc.paragraphs[:10]]).upper()
            if any(keyword in content_text for keyword in ['COMPUTER', 'TECHNOLOGY', 'PROGRAMMING']):
                subject = "Computer Science"
            elif any(keyword in content_text for keyword in ['MATH', 'ALGEBRA', 'GEOMETRY']):
                subject = "Mathematics"
            elif any(keyword in content_text for keyword in ['SCIENCE', 'PHYSICS', 'CHEMISTRY']):
                subject = "Science"
            
            return {
                'title': title,
                'subject': subject
            }
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {
                'title': "Untitled Document",
                'subject': "General"
            }
    
    def extract_text_from_txt(self, file_path):
        """
        Extract text content from a plain text file
        
        Args:
            file_path (str): Path to the .txt file
            
        Returns:
            list: List of tuples (page_number, content)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Split content into chunks (treating as pages)
            # Split by double newlines or every 1000 characters
            pages = []
            chunks = content.split('\n\n')
            
            current_page = 1
            current_chunk = ""
            
            for chunk in chunks:
                if len(current_chunk + chunk) > 1000:  # Max characters per page
                    if current_chunk:
                        pages.append((current_page, current_chunk.strip()))
                        current_page += 1
                        current_chunk = chunk
                    else:
                        pages.append((current_page, chunk.strip()))
                        current_page += 1
                        current_chunk = ""
                else:
                    current_chunk += "\n\n" + chunk if current_chunk else chunk
            
            # Add remaining content
            if current_chunk:
                pages.append((current_page, current_chunk.strip()))
            
            return pages
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT file: {str(e)}")
            return []
    
    def extract_text_from_image(self, file_path):
        """
        Extract text content from an image file using OCR
        
        Args:
            file_path (str): Path to the image file
            
        Returns:
            list: List of tuples (page_number, content)
        """
        try:
            import os
            filename = os.path.basename(file_path)
            
            placeholder_text = f"""
            Image uploaded: {filename}
            
            Note: This is an image file. Please describe the content of the image or ask specific questions about what you see.
            
            To get the best help:
            1. Describe what type of homework this is (math problems, reading comprehension, etc.)
            2. Tell me what specific questions you need help with
            3. Type out any text or problems you can see in the image
            
            I'm ready to help you with your homework once you provide these details!
            """
            
            return [(1, placeholder_text.strip())]
            
        except Exception as e:
            logger.error(f"Error processing image file: {str(e)}")
            return [(1, f"Error processing image: {str(e)}")]
